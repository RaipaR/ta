"""Document generation helpers built on top of Microsoft Word templates."""
from __future__ import annotations

from dataclasses import asdict
from datetime import date
from pathlib import Path
from typing import Dict, Mapping

from docx import Document

from .models import Booking, Tourist


def _normalise_context(context: Mapping[str, object]) -> Dict[str, str]:
    normalised: Dict[str, str] = {}
    for key, value in context.items():
        if isinstance(value, date):
            normalised[key] = value.strftime("%d.%m.%Y")
        elif value is None:
            normalised[key] = ""
        else:
            normalised[key] = str(value)
    return normalised


def _replace_text_in_paragraph(paragraph, context: Mapping[str, str]) -> None:
    for key, value in context.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)


def _replace_text_in_table(table, context: Mapping[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(paragraph, context)


def fill_template(template_path: Path, output_path: Path, context: Mapping[str, object]) -> Path:
    """Fill a Word template with provided context values."""
    document = Document(template_path)
    normalised_context = _normalise_context(context)

    for paragraph in document.paragraphs:
        _replace_text_in_paragraph(paragraph, normalised_context)

    for table in document.tables:
        _replace_text_in_table(table, normalised_context)

    document.save(output_path)
    return output_path


def render_booking_contract(
    tourist: Tourist,
    booking: Booking,
    template_path: Path,
    output_path: Path,
) -> Path:
    """Render a booking contract using a tourist, booking and template."""
    context: Dict[str, object] = {
        **{f"tourist_{key}": value for key, value in asdict(tourist).items()},
        **{f"booking_{key}": value for key, value in asdict(booking).items()},
    }
    return fill_template(template_path, output_path, context)
